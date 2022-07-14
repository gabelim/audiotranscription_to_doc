
import json
from itertools import groupby  
import docx
from operator import itemgetter

# Opening JSON file
f = open('t.json')
  
# returns JSON object as 
# a dictionary
data = json.load(f)

text=data['text']
text_list=text.split()
print(text_list[0])

#reading from "words" key in json
w=data['words']
# print(w['speaker'])

#speaker_list is huge list of speaker per word in order
speaker_list=[]
for i in w:
    # print(i)
    speaker_list.append(i['speaker'])

#speaker_order compresses consecutive speakers in speaker_list
speaker_order=[key for key, _group in groupby(speaker_list)]

#range is index range of consecutive words by a speaker 
ranges=[list(g) for _,g in groupby(range(len(speaker_list)),lambda idx:speaker_list[idx])]
#list of start and end of words spoken by a speaker
final=[[r[0],r[-1]] if len(r)>1 else r for r in ranges]



#write to word document
doc = docx.Document()
for s in range(len(speaker_order)):
    # print(s)
    # print(speaker_order[s])
    # print(len(final[s]))

    #if there is a start and end list [start:end], len is 2
    if len(final[s])==2:
        text_block=" ".join(text_list[final[s][0]:final[s][1]])
    #else it is single word spoken by one speaker
    else:
        text_block=text_list[final[s][0]]
    print(text_block)

#print speaker in order, then the related text
    doc.add_paragraph(speaker_order[s])
    doc.add_paragraph(text_block)

    doc.save('out.docx')








